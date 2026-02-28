import io
from django.http import HttpResponse
from django.template.loader import get_template
from django.utils.translation import gettext as _
from django.utils import timezone
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def check_export_permissions(user, timetable):
    """
    Check if user has permission to export this specific timetable.
    Returns True if allowed, False otherwise.
    """
    # Admin can export everything
    if user.is_staff or user.is_superuser or user.is_admin():
        return True
        
    # Students can only export their own program's timetable
    if user.is_student():
        if hasattr(user, 'student_profile') and user.student_profile.group:
            return user.student_profile.group.program == timetable.program
    
    # Teachers can export any timetable matching their sessions, 
    # but we will likely filter the content later. 
    # For now, allow access if they are a teacher, content filtering happens in view.
    if user.is_teacher():
        return True
        
    # Associations similar to teachers (or maybe restricted? sticking to plan)
    if user.is_association():
        return True
        
    return False

def generate_pdf(context, template_path='scheduling/export/pdf_template.html'):
    """
    Generate PDF from Django template.
    """
    template = get_template(template_path)
    html = template.render(context)
    result = io.BytesIO()
    
    # pisa.CreatePDF(io.BytesIO(html.encode("UTF-8")), dest=result)
    # Using specific encodings if necessary, but UTF-8 usually works
    try:
        pisa_status = pisa.CreatePDF(html, dest=result)
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None
        
    if pisa_status.err:
        return None
        
    return result.getvalue()

def set_cell_background(cell, color_hex):
    """
    Helper to set cell background color in docx.
    color_hex: string like 'FFFFFF' (no #) or 'auto'
    """
    if not color_hex or color_hex.startswith('#'):
        color_hex = color_hex[1:] if color_hex else 'FFFFFF'
        
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_docx(timetable, entries, time_slots, days):
    """
    Generate Word document for the timetable.
    """
    document = Document()
    
    # Set orientation to landscape (hacky in python-docx but workable)
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Title
    title = document.add_heading(level=1)
    run = title.add_run(f"{timetable.name}")
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{timetable.program.name} - {timetable.get_semester_display()} ({timetable.academic_year})").bold = True
    
    document.add_paragraph() # Spacer
    
    # Create Table
    # Rows: Header + Time Slots
    # Cols: Time Column + Days
    
    # Pre-process data into a grid for easier writing
    # grid[slot_index][day_index] = [entries]
    
    col_days = [d[0] for d in days] # ['MON', 'TUE'...]
    
    table = document.add_table(rows=1, cols=len(days) + 1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = _("Horaire")
    set_cell_background(hdr_cells[0], 'E0E0E0')
    
    for i, (day_code, day_name) in enumerate(days):
        cell = hdr_cells[i+1]
        cell.text = day_name
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'E0E0E0')
    
    # Organize entries by slot and day
    # entry_map = { slot_id: { day_code: [entries] } }
    entry_map = {}
    for slot in time_slots:
        entry_map[slot.id] = { day_code: [] for day_code in col_days }
        
    for entry in entries:
        if entry.time_slot_id in entry_map and entry.day_of_week in entry_map[entry.time_slot_id]:
            entry_map[entry.time_slot_id][entry.day_of_week].append(entry)
            
    # Rows
    for slot in time_slots:
        row = table.add_row()
        cells = row.cells
        
        # Time Column
        time_text = f"{slot.start_time.strftime('%H:%M')}\n-\n{slot.end_time.strftime('%H:%M')}"
        cells[0].text = time_text
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cells[0], 'F5F5F5')
        
        # Day Columns
        for i, day_code in enumerate(col_days):
            cell = cells[i+1]
            cell_entries = entry_map[slot.id][day_code]
            
            if cell_entries:
                # Add content
                for idx, entry in enumerate(cell_entries):
                    if idx > 0:
                        cell.add_paragraph("---")
                    
                    # Subject
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(entry.subject.name)
                    run.bold = True
                    run.font.size = Pt(10)
                    
                    # Type
                    p2 = cell.add_paragraph(entry.get_session_type_display())
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p2.style = document.styles['No Spacing']
                    p2.runs[0].font.size = Pt(8)
                    p2.runs[0].font.italic = True
                    
                    # Info (Room / Teacher)
                    info_text = ""
                    if entry.room:
                        info_text += f"{entry.room.name}"
                    if entry.teacher:
                        if info_text: info_text += "\n"
                        info_text += f"{entry.teacher.user.last_name}"
                    
                    p3 = cell.add_paragraph(info_text)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = document.styles['No Spacing']
                    p3.runs[0].font.size = Pt(9)
                
                # Simple color coding based on first entry type (limited in docx without complexity)
                first_type = cell_entries[0].session_type
                if first_type == 'cours':
                    set_cell_background(cell, 'DBEAFE') # Light Blue
                elif first_type == 'td':
                    set_cell_background(cell, 'D1FAE5') # Light Green
                elif first_type == 'tp':
                    set_cell_background(cell, 'FEF3C7') # Light Orange
            else:
                set_cell_background(cell, 'FFFFFF')

    # Add Footer with timestamp
    document.add_paragraph()
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    timestamp = timezone.now().strftime("%d/%m/%Y %H:%M")
    footer.add_run(f"Généré par FSTTIME le {timestamp}").font.size = Pt(8)
    
    # Save to IO
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f
